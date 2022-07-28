from docx import Document
document = Document('demo.docx')
choice = int(input('1.Read file\n2.Write text to file\n'))
if choice == 1:
    for i in range(len(document.paragraphs)):
        print(document.paragraphs[i].text)
elif choice == 2:
    document.add_paragraph(input())
    document.save('demo.docx')
else:
    print("Unknown command")